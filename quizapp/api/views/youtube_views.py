from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from ..chains.youtube import generate, YouTubeContentGenerator
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
import logging
from rest_framework.decorators import api_view, permission_classes
from ..models.youtube_models import YouTubeContent

logger = logging.getLogger(__name__)

# Create a singleton instance of YouTubeContentGenerator
youtube_generator = YouTubeContentGenerator()

class YouTubeProcessView(APIView):
    permission_classes = [IsAuthenticated]

    def extract_video_id(self, url):
        # Regular expressions for different YouTube URL formats
        patterns = [
            r'(?:youtube\.com\/watch\?v=|youtu\.be\/)([^&\n?]+)',
            r'youtube\.com\/embed\/([^&\n?]+)',
            r'youtube\.com\/v\/([^&\n?]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        return None

    def post(self, request, format=None):
        url = request.data.get('url')
        if not url:
            return Response({'error': 'URL is required'}, status=status.HTTP_400_BAD_REQUEST)

        video_id = self.extract_video_id(url)
        if not video_id:
            return Response({'error': 'Invalid YouTube URL'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Process the video based on the endpoint
            endpoint = request.path.split('/')[-2]  # Get the endpoint (quiz, flashcards, notes, or chat)
            
            # Get number of questions for quiz generation
            num_questions = request.data.get('num_questions', 5)  # Default to 5 questions
            
            # Call the Gemini API to process the video
            result = generate(video_id, endpoint, num_questions=num_questions)
            
            # For chat endpoint, initialize the chat with the video content
            if endpoint == 'chat':
                # Process the video for chat
                chat_response = youtube_generator.process_video_for_chat(video_id)
                return Response({
                    'id': video_id,
                    'title': 'Video Chat',
                    'content': chat_response,
                    'video_url': f"https://youtu.be/{video_id}"
                })
            
            # Save the content to the database
            content_type = endpoint
            if content_type == 'quizz':  # Fix the content type name
                content_type = 'quiz'
            
            # Create YouTubeContent instance
            youtube_content = YouTubeContent.objects.create(
                title=result.get('title', 'Video Content'),
                video_url=f"https://youtu.be/{video_id}",
                thumbnail_url=f"https://img.youtube.com/vi/{video_id}/maxresdefault.jpg",
                content_type=content_type,
                content_data=result.get('content', {}),
                user=request.user
            )
            
            # Return the generated content
            if isinstance(result, dict):
                return Response({
                    'id': youtube_content.id,
                    'title': result.get('title', ''),
                    'content': result.get('content', {}),
                    'video_url': f"https://youtu.be/{video_id}"
                })
            else:
                return Response({
                    'id': youtube_content.id,
                    'title': 'Video Content',
                    'content': result,
                    'video_url': f"https://youtu.be/{video_id}"
                })
        except Exception as e:
            logger.error(f"Error in YouTubeProcessView: {str(e)}", exc_info=True)
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class YouTubeNotesDownloadView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        try:
            notes = request.data.get('notes', [])
            title = request.data.get('title', 'Video Notes')
            video_url = request.data.get('video_url', '')

            # Create a new Word document
            doc = Document()

            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add video URL
            if video_url:
                url_paragraph = doc.add_paragraph()
                url_run = url_paragraph.add_run(f"Source: {video_url}")
                url_run.italic = True
                url_run.font.size = Pt(10)
                url_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            timestamp_paragraph = doc.add_paragraph()
            timestamp_run = timestamp_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_run.italic = True
            timestamp_run.font.size = Pt(10)
            timestamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add a line break
            doc.add_paragraph()

            # Add each section
            for section in notes:
                # Add section title
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(section['title'])
                title_run.bold = True
                title_run.font.size = Pt(14)

                # Add section content
                content_paragraph = doc.add_paragraph()
                content_run = content_paragraph.add_run(section['content'])
                content_run.font.size = Pt(12)

                # Add spacing between sections
                doc.add_paragraph()

            # Save the document to a bytes buffer
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)

            # Create the response
            response = Response(
                docx_bytes.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{title.replace(" ", "_")}_notes.docx"'
            
            return response

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

class YouTubeChatMessageView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        try:
            message = request.data.get('message')
            video_id = request.data.get('video_id')
            
            logger.info(f"Received chat message request - Message: {message}, Video ID: {video_id}")
            
            if not message:
                logger.warning("Missing message in request")
                return Response(
                    {'error': 'Message is required'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            if not video_id:
                logger.warning("Missing video_id in request")
                return Response(
                    {'error': 'Video ID is required'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Use the singleton instance
            response = youtube_generator.handle_chat_message(message, video_id)
            
            logger.info(f"Successfully generated chat response for video {video_id}")
            return Response({
                'content': response
            })
            
        except Exception as e:
            logger.error(f"Error in YouTubeChatMessageView: {str(e)}", exc_info=True)
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def youtube_history(request):
    try:
        # Get the latest 4 YouTube content items for the current user
        content = YouTubeContent.objects.filter(user=request.user).order_by('-created_at')[:4]
        
        # Convert to list of dictionaries
        content_list = [{
            'id': item.id,
            'title': item.title,
            'video_url': item.video_url,
            'thumbnail_url': item.thumbnail_url,
            'content_type': item.content_type,
            'content_data': item.content_data,
            'created_at': item.created_at.isoformat(),
            'user_id': item.user_id
        } for item in content]
        
        return Response(content_list)
    except Exception as e:
        logger.error(f"Error in youtube_history: {str(e)}", exc_info=True)
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR) 